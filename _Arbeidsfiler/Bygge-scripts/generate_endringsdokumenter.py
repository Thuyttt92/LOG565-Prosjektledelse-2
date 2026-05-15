"""
Genererer to endringsdokumenter for LOG565 — Nye Hædda Barneskole:
  1. NHB-IRGESUND — Schedule crashing av 4.1 Råbygg (kommunestyrets vedtak 07.05.2026)
  2. CR-001 — Oppgradert sprinkler/rømning på 5.1 VVS (DSB-veileder, mnd 13)

Bygger fra malen 03 - Gjennomføring/Maler/endringsdokument_mal.docx
"""
from __future__ import annotations
import shutil
from pathlib import Path
from docx import Document

import sys
sys.path.insert(0, str(Path(__file__).resolve().parent))
from paths import ROOT, MAL_ENDRINGSDOKUMENT as MAL, ENDRINGSDOK as OUTDIR
OUTDIR.mkdir(exist_ok=True)


def set_cell(cell, *lines: str) -> None:
    """Replace a cell's text with one or more paragraphs."""
    # Clear all existing paragraphs except the first
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    # Set first paragraph text
    cell.paragraphs[0].clear()
    cell.paragraphs[0].add_run(lines[0])
    # Add remaining as new paragraphs
    for line in lines[1:]:
        cell.add_paragraph(line)


def fill_header_table(doc, *, prosjekt, endrings_id, dato, versjon,
                      tittel, initiert_av, vurderer, status,
                      forum, prioritet, beslutningsdato, vindu):
    # Table 1 = header (prosjekt-info), see mal-struktur
    t = doc.tables[1]
    # Row 0
    set_cell(t.rows[0].cells[0], "Prosjekt", prosjekt)
    set_cell(t.rows[0].cells[1], "Endrings-ID", endrings_id)
    set_cell(t.rows[0].cells[2], "Dato", dato)
    set_cell(t.rows[0].cells[3], "Versjon", versjon)
    # Row 1
    set_cell(t.rows[1].cells[0], "Endringstittel", tittel)
    set_cell(t.rows[1].cells[1], "Initiert av", initiert_av)
    set_cell(t.rows[1].cells[2], "Ansvarlig vurderer", vurderer)
    set_cell(t.rows[1].cells[3], "Status", status)
    # Row 2
    set_cell(t.rows[2].cells[0], "Beslutningsforum", forum)
    set_cell(t.rows[2].cells[1], "Prioritet", prioritet)
    set_cell(t.rows[2].cells[2], "Ønsket beslutningsdato", beslutningsdato)
    set_cell(t.rows[2].cells[3], "Implementeringsvindu", vindu)


def fill_beskrivelse(doc, *, bakgrunn, foreslatt, begrunnelse, konsekvens_uten):
    # Table 3
    t = doc.tables[3]
    set_cell(t.rows[0].cells[1], bakgrunn)
    set_cell(t.rows[1].cells[1], foreslatt)
    set_cell(t.rows[2].cells[1], begrunnelse)
    set_cell(t.rows[3].cells[1], konsekvens_uten)


def fill_konsekvens(doc, rows):
    """rows = list of 6 tuples: (paavirkning, beskrivelse, kommentar) for
    Omfang, Tid, Kostnad, Kvalitet, Ressurser, Risiko/HMS."""
    t = doc.tables[4]
    for i, (paavirk, beskr, komm) in enumerate(rows, start=1):
        set_cell(t.rows[i].cells[1], paavirk)
        set_cell(t.rows[i].cells[2], beskr)
        set_cell(t.rows[i].cells[3], komm)


def fill_styringsrammer(doc, *, sluttdato, budsjett, risikobudsjett, tidsbuffer,
                         eac, prognose_sluttdato, kontraktsmessig, rebaseline):
    t = doc.tables[5]
    set_cell(t.rows[0].cells[0], "Endring i sluttdato", sluttdato)
    set_cell(t.rows[0].cells[1], "Endring i budsjett", budsjett)
    set_cell(t.rows[0].cells[2], "Bruk av risikobudsjett", risikobudsjett)
    set_cell(t.rows[0].cells[3], "Bruk av tidsbuffer", tidsbuffer)
    set_cell(t.rows[1].cells[0], "Ny prognose (EAC)", eac)
    set_cell(t.rows[1].cells[1], "Ny prognose sluttdato", prognose_sluttdato)
    set_cell(t.rows[1].cells[2], "Kontraktsmessig konsekvens", kontraktsmessig)
    set_cell(t.rows[1].cells[3], "Behov for re-baselining", rebaseline)


def fill_beslutning(doc, *, anbefaling, beslutning, begrunnelse, forutsetninger):
    t = doc.tables[6]
    set_cell(t.rows[0].cells[1], anbefaling)
    set_cell(t.rows[1].cells[1], beslutning)
    set_cell(t.rows[2].cells[1], begrunnelse)
    set_cell(t.rows[3].cells[1], forutsetninger)


def fill_implementering(doc, *, tiltak, ansvar, frist, styringsdok, lukking):
    t = doc.tables[7]
    set_cell(t.rows[0].cells[1], tiltak)
    set_cell(t.rows[1].cells[1], ansvar)
    set_cell(t.rows[2].cells[1], frist)
    set_cell(t.rows[3].cells[1], styringsdok)
    set_cell(t.rows[4].cells[1], lukking)


def fill_signatur(doc, *, foreslatt_av, vurdert_av, godkjent_av,
                   dato_foreslatt, dato_vurdert, dato_godkjent):
    t = doc.tables[8]
    set_cell(t.rows[1].cells[1], foreslatt_av)
    set_cell(t.rows[1].cells[2], dato_foreslatt)
    set_cell(t.rows[2].cells[1], vurdert_av)
    set_cell(t.rows[2].cells[2], dato_vurdert)
    set_cell(t.rows[3].cells[1], godkjent_av)
    set_cell(t.rows[3].cells[2], dato_godkjent)


def fill_historikk(doc, *, dato, endring, utarbeidet):
    t = doc.tables[9]
    set_cell(t.rows[1].cells[1], dato)
    set_cell(t.rows[1].cells[2], endring)
    set_cell(t.rows[1].cells[3], utarbeidet)


# =============================================================================
# DOKUMENT 1: NHB-IRGESUND — Schedule crashing
# =============================================================================
def lag_nhb_irgesund():
    out = OUTDIR / "Endringsdokument NHB-IRGESUND - Schedule crashing.docx"
    shutil.copy(MAL, out)
    doc = Document(out)

    fill_header_table(
        doc,
        prosjekt="Nye Hædda Barneskole — LOG565",
        endrings_id="NHB-IRGESUND",
        dato="07.05.2026",
        versjon="1.0",
        tittel="Komprimering (crashing) av kritisk aktivitet 4.1 Råbygg og utvidelse av budsjettramme fra 700 til 800 MNOK",
        initiert_av="Prosjektledelsen, Nye Hædda Barneskole",
        vurderer="Byggekomiteen",
        status="Godkjent",
        forum="Kommunestyret (Hædda kommune)",
        prioritet="Høy",
        beslutningsdato="07.05.2026",
        vindu="Umiddelbart etter vedtak — innarbeides i Baseline 1",
    )

    fill_beskrivelse(
        doc,
        bakgrunn=(
            "Detaljprosjektering har avdekket at aktivitet 4.1 Råbygg er underestimert i den opprinnelige planen. "
            "Aktivitetens reelle behov er +50 MNOK og +2 måneder ut over det opprinnelige estimatet — varigheten økes "
            "fra 5 til 7 måneder, og kostnaden økes fra 140 til 190 MNOK. Aktiviteten ligger på prosjektets kritiske vei, "
            "slik at overskridelsen forplanter seg direkte til prosjektets totalkostnad og sluttdato. "
            "Nåværende prognose uten korrigerende tiltak: totalkostnad 750 MNOK (+50 MNOK), sluttdato juli 2026 (forsinkelse ~2 mnd "
            "mot vedtatt sluttdato 15. mai 2026). Tid er definert som den viktigste rammebetingelsen — skolen skal stå klar til "
            "skolestart høsten 2026, og en forsinkelse vil utløse betydelige følgekostnader (midlertidige skolelokaler, "
            "ny oppstartsplanlegging, omdømmebelastning)."
        ),
        foreslatt=(
            "Komprimere aktivitet 4.1 Råbygg ved å øke ressurspådraget gjennom overtid, parallelle skift, premium-leveranser "
            "av materialer og forsering hos underleverandører. Tiltaket retter seg utelukkende mot prosjektets kritiske vei "
            "og bringer varigheten på 4.1 Råbygg tilbake fra 7 til 5 måneder. Samtidig utvides budsjettrammen fra 700 MNOK "
            "til 800 MNOK for å finansiere både overskridelsen (50 MNOK) og kostnaden ved selve komprimeringen (50 MNOK)."
        ),
        begrunnelse=(
            "Vedtatt sluttdato 15. mai 2026 overholdes fullt ut, slik at skolestart høsten 2026 er sikret. Risikoen er "
            "konsentrert til én kjent aktivitet på den kritiske veien og lar seg styre med tett oppfølging. Skolestart-løftet "
            "til kommunens innbyggere holdes, og betydelige følgekostnader unngås. Komprimering er det eneste tiltaket "
            "innenfor prosjektets handlingsrom som kan bringe sluttdatoen tilbake til den vedtatte fristen."
        ),
        konsekvens_uten=(
            "Prosjektet leveres juli 2026 til en kostnad på 750 MNOK. Tidsrammen brytes med ca. 2 måneder og "
            "kostnadsrammen brytes med 50 MNOK — begge styringsrammer brytes samtidig. Skolestart høsten 2026 må "
            "utsettes eller gjennomføres i midlertidige lokaler, med tilhørende omdømmebelastning for kommunen."
        ),
    )

    fill_konsekvens(doc, [
        ("Ingen endring",
         "Leveransene er de samme — kun varighet og kostnad på 4.1 Råbygg endres. Kravspesifikasjonen leveres uendret.",
         "Bekreft at premium-leveranser og forsering ikke medfører kvalitetsreduksjon."),
        ("Reduksjon (–2 mnd på kritisk vei)",
         "Varigheten på 4.1 Råbygg reduseres fra 7 til 5 måneder. Vedtatt sluttdato 15. mai 2026 overholdes.",
         "Effekt på kritisk vei verifiseres ved re-baselining i MS Project."),
        ("Økning (+100 MNOK på totalramme)",
         "Budsjettrammen utvides fra 700 MNOK til 800 MNOK: 50 MNOK dekker den avdekkede overskridelsen på 4.1, "
         "50 MNOK dekker selve komprimeringen.",
         "Risikoreserven på 50 MNOK forblir intakt og benyttes ikke til crashing."),
        ("Lav risiko — krever skjerpet kontroll",
         "Premium-materialer og forsering kan øke belastning på leverandører. Kvalitetskontroll skjerpes på 4.1 Råbygg.",
         "Følges opp via problemliste og månedsrapport gjennom hele 4.1-perioden (mnd 7–11)."),
        ("Økt belastning på underleverandører",
         "Forsering forutsetter at underleverandører kan levere parallelt skift og premium-leveranser. Avtaler inngås "
         "med tilleggsbestilling.",
         "Kontraktsmessige tillegg signeres før crashing iverksettes."),
        ("Økt HMS-risiko ved forsering",
         "Parallelle skift og overtid øker risiko for HMS-hendelser. SHA-oppfølging styrkes på byggeplassen.",
         "Ukentlig SHA-statusmøte etableres for 4.1-perioden."),
    ])

    fill_styringsrammer(
        doc,
        sluttdato="Ingen (15.05.2026 opprettholdes)",
        budsjett="+100 MNOK (700 → 800 MNOK)",
        risikobudsjett="0 MNOK (risikoreserve 50 MNOK forblir uendret og separat)",
        tidsbuffer="0 uker (8 ukers tidsbuffer forblir uendret og separat)",
        eac="800 MNOK",
        prognose_sluttdato="15.05.2026",
        kontraktsmessig="Ja — tilleggsbestillinger til underleverandører på 4.1 Råbygg (premium-leveranser, forsering, overtid)",
        rebaseline="Ja — Baseline 1 settes i MS Project umiddelbart etter kommunestyrets vedtak",
    )

    fill_beslutning(
        doc,
        anbefaling="Anbefales godkjent.",
        beslutning="Godkjent av kommunestyret 07.05.2026 (sak nr. NHB-IRGESUND).",
        begrunnelse=(
            "Tid er definert som den viktigste rammebetingelsen for prosjektet — skolen skal stå klar til skolestart "
            "høsten 2026, og tidsrammen kan i praksis ikke forskyves. Vurderte alternativer (ikke iverksette tiltak, "
            "redusere omfang) er forkastet: «ikke iverksette tiltak» bryter både tids- og kostnadsramme samtidig, mens "
            "«redusere omfang» ikke er forenelig med kravspesifikasjonen og kommunestyrets opprinnelige bestilling. "
            "Komprimering overholder tidsrammen fullt ut mot at den økonomiske rammen utvides med 100 MNOK, og risikoen "
            "er konsentrert til én kjent aktivitet på den kritiske veien."
        ),
        forutsetninger=(
            "1) Prosjektledelsen rapporterer status på den komprimerte aktiviteten 4.1 Råbygg månedlig til byggekomiteen "
            "frem til aktiviteten er ferdigstilt. "
            "2) Kvalitetskontroll skjerpes på 4.1 Råbygg for å unngå at forsering kompromitterer leveransekvaliteten. "
            "3) SHA-oppfølging styrkes med ukentlig statusmøte i hele 4.1-perioden."
        ),
    )

    fill_implementering(
        doc,
        tiltak=(
            "1) Innhent fremskyndede tilbud fra underleverandører på 4.1 Råbygg. "
            "2) Oppdater WBS og MS Project med komprimert varighet (5 mnd) og økt kostnad (190 MNOK på 4.1). "
            "3) Sett Baseline 1 («Godkjent baseline etter crashing») i MS Project. Bevar Baseline 0 («Opprinnelig estimat»). "
            "4) Oppdater komplett prosjektplan, kostnadsprognose og risikoregister. "
            "5) Etabler månedlig statusrapportering til byggekomiteen om 4.1 Råbygg. "
            "6) Kommuniser ny totalramme (800 MNOK) og crashing-vedtaket til alle teamledere."
        ),
        ansvar="Prosjektleder (gruppe irgesundinger)",
        frist="Implementering ferdigstilt før crashing-perioden starter (planlagt oppstart 4.1 Råbygg: august 2025)",
        styringsdok=(
            "Komplett prosjektplan (kap. om endringsstyring og rammer), MS Project (Baseline 1), risikoregister "
            "(R-006 leverandøroverbelastning under forsering), kostnadsprognose, kontraktsdokumenter med underleverandører."
        ),
        lukking=(
            "Endringen verifiseres lukket når 4.1 Råbygg er ferdigstilt iht. komprimert plan (planlagt desember 2025) "
            "og prosjektet er levert innenfor 15.05.2026. Endelig dokumentasjon i sluttrapport (fase 4)."
        ),
    )

    fill_signatur(
        doc,
        foreslatt_av="Prosjektledelsen, Nye Hædda Barneskole",
        dato_foreslatt="07.05.2026",
        vurdert_av="Byggekomiteen",
        dato_vurdert="07.05.2026",
        godkjent_av="Kommunestyret (Hædda kommune)",
        dato_godkjent="07.05.2026",
    )

    fill_historikk(
        doc,
        dato="07.05.2026",
        endring="Første utkast — opprinnelig registrering av kommunestyrets vedtak NHB-IRGESUND",
        utarbeidet="Prosjektkoordinator",
    )

    doc.save(out)
    print(f"  -> {out.name}")


# =============================================================================
# DOKUMENT 2: CR-001 — Sprinkler-rømning
# =============================================================================
def lag_cr_001_sprinkler():
    out = OUTDIR / "Endringsdokument CR-001 - Sprinkler-romning.docx"
    shutil.copy(MAL, out)
    doc = Document(out)

    fill_header_table(
        doc,
        prosjekt="Nye Hædda Barneskole — LOG565",
        endrings_id="CR-001",
        dato="03.02.2026",
        versjon="1.0",
        tittel="Oppgradert sprinklerdekning og rømningsskilting på 5.1 VVS iht. ny DSB-veileder",
        initiert_av="Fagansvarlig VVS",
        vurderer="Prosjektleder",
        status="Godkjent",
        forum="Prosjekteier (Hædda kommune)",
        prioritet="Høy",
        beslutningsdato="03.02.2026",
        vindu="Innarbeides i pågående installasjonssekvens på 5.1 VVS (mnd 13–14)",
    )

    fill_beskrivelse(
        doc,
        bakgrunn=(
            "Direktoratet for samfunnssikkerhet og beredskap (DSB) har publisert oppdatert veileder for sprinklerdekning "
            "og rømningsskilting i kommunale byggeprosjekter. Veilederen trer i kraft før prosjektets ferdigstillelse "
            "(15.05.2026). Endringen påvirker arbeidspakke 5.1 VVS og kan ikke ignoreres uten å risikere at skolen "
            "ikke får brukstillatelse fra kommunen før skolestart høsten 2026. Saken ble meldt i teamledermøtet "
            "03.02.2026 og registrert som risiko R-06 i risikoregisteret."
        ),
        foreslatt=(
            "Oppgradere sprinklerdekning og forsterke rømningsskilting i hele bygningsmassen iht. ny DSB-veileder. "
            "Omfanget av leveranse 5.1 VVS utvides tilsvarende. Eksisterende VVS-entreprenør tar tilleggsbestilling og "
            "innarbeider endringen i pågående installasjonssekvens. Etterfølgende innvendige arbeider replanlegges "
            "lokalt for å begrense kaskadeeffekt."
        ),
        begrunnelse=(
            "Sikrer brukstillatelse fra kommunen og etterlevelse av offentlige krav, slik at skolestart høsten 2026 "
            "ikke trues. Sikkerhetsstandard for elever og ansatte heves. Forhindrer kostbar og forstyrrende "
            "etterinstallasjon etter at skolen er tatt i bruk."
        ),
        konsekvens_uten=(
            "Skolen risikerer å ikke få brukstillatelse ved planlagt overlevering. Skolestart høsten 2026 må enten "
            "utsettes eller gjennomføres i midlertidige lokaler, med betydelige følgekostnader og omdømmebelastning."
        ),
    )

    fill_konsekvens(doc, [
        ("Middels — scope-utvidelse på 5.1",
         "Sprinklerdekning utvides og rømningsskilting forsterkes iht. ny DSB-veileder. Øvrige leveranser uendret.",
         "Tilleggsleveranse formaliseres i kontraktstillegg med VVS-entreprenør."),
        ("Lav — 1 uke forskyvning",
         "Installasjonssekvensen på 5.1 VVS får ca. 1 uke forskyvning. Absorberes innenfor godkjent tidsbuffer (8 uker).",
         "Kumulativ tidsbuffer-bruk etter mnd 13: 1,5 uker av 8 uker."),
        ("Middels — 5,0 MNOK fra risikoreserve",
         "Tilleggskostnaden 5,0 MNOK dekkes av godkjent risikoreserve (50 MNOK). BAC på 800 MNOK påvirkes ikke.",
         "Kumulativ risikoreserve-bruk etter mnd 13: 11,0 MNOK av 50 MNOK."),
        ("Forbedring",
         "Strengere sikkerhetsstandard for sprinkler og rømning gir bedre brann- og rømningssikkerhet for elever og ansatte.",
         "Godkjennes av brannvesen ved sluttbefaring (BP3)."),
        ("Lav — eksisterende leverandør",
         "Eksisterende VVS-entreprenør utvider leveransen. Ingen behov for ny anbudsrunde.",
         "Tilleggsbestilling håndteres innenfor rammeavtale."),
        ("Forbedring",
         "Bedre brann- og rømningssikkerhet reduserer HMS-risiko under drift av skolen.",
         "Risiko R-06 lukkes når 5.1 VVS er ferdigstilt og DSB-krav er etterlevd."),
    ])

    fill_styringsrammer(
        doc,
        sluttdato="Ingen (15.05.2026 opprettholdes — forsinkelse absorbert i tidsbuffer)",
        budsjett="Ingen endring i BAC (5,0 MNOK dekkes av risikoreserve)",
        risikobudsjett="5,0 MNOK (av godkjent 50,0 MNOK)",
        tidsbuffer="1 uke (av godkjent 8 uker)",
        eac="800 MNOK (uendret)",
        prognose_sluttdato="15.05.2026 (uendret)",
        kontraktsmessig="Ja — tilleggsbestilling til VVS-entreprenør på 5.1",
        rebaseline="Nei — endring håndteres innenfor godkjent risikoreserve og tidsbuffer; baseline opprettholdes",
    )

    fill_beslutning(
        doc,
        anbefaling="Anbefales godkjent.",
        beslutning="Godkjent av prosjekteier 03.02.2026.",
        begrunnelse=(
            "Regulatorisk krav som ikke kan ignoreres uten å risikere brukstillatelse. Kostnaden (5,0 MNOK) dekkes "
            "av godkjent risikoreserve, og forsinkelsen (1 uke) absorberes i godkjent tidsbuffer. Vedtatt sluttdato "
            "15.05.2026 og totalramme 800 MNOK opprettholdes uendret. Tiltaket forbedrer i tillegg HMS-nivået på "
            "leveransen."
        ),
        forutsetninger=(
            "1) Installasjonssekvensen på 5.1 VVS replanlegges lokalt for å minimere kaskadeeffekt på etterfølgende "
            "innvendige arbeider. "
            "2) Ukentlig oppfølging av berørte arbeidspakker frem til 5.1 VVS er ferdigstilt. "
            "3) Frist for ansvarlig oppfølging: 14.03.2026 (prosjektleder)."
        ),
    )

    fill_implementering(
        doc,
        tiltak=(
            "1) Inngå kontraktstillegg med VVS-entreprenør på 5.1 (oppgradert sprinkler og rømning). "
            "2) Oppdater installasjonsplan og bemanningsskjema for 5.1 VVS. "
            "3) Replanlegg etterfølgende innvendige arbeider (4.2/4.3/4.4) for å begrense kaskadeeffekt. "
            "4) Oppdater risikoregister (R-06 oppgraderes med ny status). "
            "5) Oppdater kostnadsprognose med 5,0 MNOK uttrekk fra risikoreserve. "
            "6) Oppdater MS Project-fremdrift med 1 uke forskyvning på 5.1."
        ),
        ansvar="Prosjektleder (gruppe irgesundinger)",
        frist="14.03.2026 (lukking av ansvarlig oppfølging iht. møtereferat mnd 13)",
        styringsdok=(
            "Risikoregister (R-06), kostnadsprognose, MS Project (fremdrift), kontraktsdokumenter med VVS-entreprenør, "
            "FDV-dokumentasjon (oppdaterte sprinkler- og rømningstegninger)."
        ),
        lukking=(
            "Endringen verifiseres lukket når 5.1 VVS er ferdigstilt iht. ny DSB-veileder (ferdigstilt mnd 15, april 2026) "
            "og brannvesenet har godkjent leveransen ved ferdigbefaring (BP3, mnd 15)."
        ),
    )

    fill_signatur(
        doc,
        foreslatt_av="Fagansvarlig VVS",
        dato_foreslatt="03.02.2026",
        vurdert_av="Prosjektleder, gruppe irgesundinger",
        dato_vurdert="03.02.2026",
        godkjent_av="Prosjekteier (Hædda kommune)",
        dato_godkjent="03.02.2026",
    )

    fill_historikk(
        doc,
        dato="03.02.2026",
        endring="Første utkast — opprinnelig registrering i teamledermøte mnd 13",
        utarbeidet="Prosjektkoordinator",
    )

    doc.save(out)
    print(f"  -> {out.name}")


def main():
    # Fjern det gamle CR-001 schedule crashing-placeholderet (numreringen er
    # nå tildelt sprinkler-saken av Bård via månedsrapport mnd 13)
    gammel = OUTDIR / "Endringsdokument CR-001 - Schedule crashing.docx"
    if gammel.exists():
        gammel.unlink()
        print(f"Fjernet utdatert placeholder: {gammel.name}")

    print("Genererer endringsdokumenter:")
    lag_nhb_irgesund()
    lag_cr_001_sprinkler()
    print("Ferdig.")


if __name__ == "__main__":
    main()
