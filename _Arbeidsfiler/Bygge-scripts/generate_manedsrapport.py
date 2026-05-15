"""
Genererer månedsrapport for en gitt måned (1–16) av gjennomføringsfasen.

Bygger fra mal: 03 - Gjennomføring/Maler/månedsrapport-mal.docx
Outputfil: 03 - Gjennomføring/Månedsrapporter/Månedsrapport mnd XX - <Måned ÅÅÅÅ>.docx

Trekker data fra log565_master_data.py og inkluderer S-kurve-PNG fra
Arbeidsfiler/s_kurver/. Sporbarhet sikres ved å referere risiko-ID,
CR-ID og kilderefer i hver tabellrad.
"""
from __future__ import annotations
import sys
import shutil
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL

sys.path.insert(0, str(Path(__file__).parent))
from log565_master_data import (
    BAC, RISIKORESERVE, TIDSBUFFER_UKER, PROSJEKTLEDER_GRUPPE, SAKSNUMMER_NHB,
    MÅNEDER, PAKKER, HENDELSER,
    FREMDRIFT_PER_MND, RISIKORESERVE_BRUKT_KUM, TIDSBUFFER_BRUKT_KUM_UKER,
    beregn_evm, hent_pct_fullført,
)

sys.path.insert(0, str(Path(__file__).resolve().parent))
from paths import ROOT, MAL_MANEDSRAPPORT as MAL, MANEDSRAPPORTER as OUTDIR, S_KURVER
OUTDIR.mkdir(exist_ok=True)


# =============================================================================
# Cell-manipulasjon
# =============================================================================
def set_cell(cell, *lines: str, bold: bool = False) -> None:
    """Erstatt celletekst med en eller flere paragrafer."""
    # Fjern alle eksisterende paragrafer unntatt den første
    for p in list(cell.paragraphs[1:]):
        p._element.getparent().remove(p._element)
    # Tøm første paragraf
    p0 = cell.paragraphs[0]
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)
    run = p0.add_run(str(lines[0]))
    if bold:
        run.bold = True
    for line in lines[1:]:
        p = cell.add_paragraph()
        run = p.add_run(str(line))
        if bold:
            run.bold = True


def kr(mnok: float) -> str:
    """Format MNOK som kr-streng."""
    return f"{mnok * 1_000_000:,.0f} kr".replace(",", " ")


def mnok_fmt(mnok: float) -> str:
    return f"{mnok:.1f} MNOK"


# =============================================================================
# Tabellbygging — utvider tabeller etter behov
# =============================================================================
def ensure_rows(table, n_rows: int) -> None:
    """Sørg for at tabellen har minst n_rows rader. Klon siste rad om nødvendig."""
    while len(table.rows) < n_rows:
        # Klon siste rad
        tr = table.rows[-1]._tr
        new_tr = deepcopy(tr)
        # Fjern eksisterende innhold i kopien
        for tc in new_tr.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc"):
            for p in list(tc.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p")):
                for r in list(p.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")):
                    p.remove(r)
        tr.addnext(new_tr)


# =============================================================================
# Tabell 1 — Metadata
# =============================================================================
def fyll_metadata(doc, måned: int) -> None:
    navn, statusdato = MÅNEDER[måned]
    evm = beregn_evm(måned)
    t = doc.tables[1]
    set_cell(t.rows[0].cells[1], "Nye Hædda Barneskole")
    set_cell(t.rows[0].cells[3], f"{navn}")
    set_cell(t.rows[1].cells[1], f"Gruppe {PROSJEKTLEDER_GRUPPE}")
    set_cell(t.rows[1].cells[3], statusdato)
    set_cell(t.rows[2].cells[1], "Hædda kommune (kommunestyret)")
    set_cell(t.rows[2].cells[3], "v1.0")
    set_cell(t.rows[3].cells[1], "Prosjektledelsen, Nye Hædda Barneskole")
    set_cell(t.rows[3].cells[3], "kr")
    set_cell(t.rows[4].cells[1], evm.status_rag,
             bold=False)
    set_cell(t.rows[4].cells[3], "Byggekomiteen, prosjekteier, byggherreombud, teamledere")


# =============================================================================
# Tabell 2 — KPI-sammendrag for hele prosjektet
# =============================================================================
def fyll_kpi_sammendrag(doc, måned: int) -> None:
    evm = beregn_evm(måned)
    t = doc.tables[2]
    # Row 1: BAC | PV kum
    set_cell(t.rows[1].cells[1], kr(BAC))
    set_cell(t.rows[1].cells[3], kr(evm.pv_kum))
    # Row 2: EV kum | AC kum
    set_cell(t.rows[2].cells[1], kr(evm.ev_kum))
    set_cell(t.rows[2].cells[3], kr(evm.ac_kum))
    # Row 3: CPI | SPI
    set_cell(t.rows[3].cells[1], f"{evm.cpi:.3f}")
    set_cell(t.rows[3].cells[3], f"{evm.spi:.3f}")
    # Row 4: EAC | ETC
    set_cell(t.rows[4].cells[1], kr(evm.eac))
    set_cell(t.rows[4].cells[3], kr(evm.etc))
    # Row 5: VAC | Godkjent risikoreserve / tidsbuffer
    set_cell(t.rows[5].cells[1], kr(evm.vac))
    set_cell(t.rows[5].cells[3], f"{kr(RISIKORESERVE)} / {TIDSBUFFER_UKER} uker")
    # Row 6: Brukt/gjenværende risikoreserve | Brukt/gjenværende tidsbuffer
    rr_rest = RISIKORESERVE - evm.risikoreserve_brukt_kum
    tb_rest = TIDSBUFFER_UKER - evm.risikoreserve_brukt_uker_kum
    set_cell(t.rows[6].cells[1],
             f"Brukt {kr(evm.risikoreserve_brukt_kum)} / Gjenværende {kr(rr_rest)}")
    set_cell(t.rows[6].cells[3],
             f"Brukt {evm.risikoreserve_brukt_uker_kum:.1f} uker / Gjenværende {tb_rest:.1f} uker")


# =============================================================================
# Tabell 3 — KPI per milepæl (WBS-nivå 1)
# =============================================================================
WBS_NIVA1 = [
    ("1", "Prosjektledelse og Administrasjon", ["1.1", "1.2", "1.3", "1.4"]),
    ("2", "Planlegging og Prosjektering", ["2.1", "2.2", "2.3"]),
    ("3", "Forberedelse og Riving", ["3.1", "3.2", "3.3"]),
    ("4", "Skolebygg – Bygningsmessige arbeider", ["4.1", "4.2", "4.3", "4.4", "4.5"]),
    ("5", "Skolebygg – Tekniske Anlegg", ["5.1", "5.2", "5.3", "5.4", "5.5"]),
    ("6", "Utomhus og Uteområder", ["6.1", "6.2", "6.3", "6.4"]),
    ("7", "Inventar og Utstyr (FF&E)", ["7.1", "7.2", "7.3"]),
    ("8", "Overtakelse og Avslutning", ["8.1", "8.2", "8.3", "8.4", "8.5"]),
]


def fyll_milepaeler(doc, måned: int) -> None:
    """Aggregert EVM per WBS-nivå 1-gruppe."""
    from log565_master_data import PAKKE_BY_WBS, beregn_pv as _beregn_pv
    t = doc.tables[3]
    # Tabellen har 7 rader (1 header + 6 milepæler). Vi trenger 9 rader (1+8).
    ensure_rows(t, 9)

    for i, (mid, mnavn, wbs_list) in enumerate(WBS_NIVA1, start=1):
        bac_g = sum(PAKKE_BY_WBS[w].bac for w in wbs_list)
        ev_g = sum(PAKKE_BY_WBS[w].bac * hent_pct_fullført(w, måned) / 100
                   for w in wbs_list)
        # AC per milepæl: estimat ut fra kumulativ påløpt for pakkene
        # (vi har ikke per-pakke AC, men kan estimere fra EV × CPI_global)
        evm_global = beregn_evm(måned)
        ac_g = ev_g / evm_global.cpi if evm_global.cpi > 0 else 0
        cpi_g = ev_g / ac_g if ac_g > 0 else 0
        # PV per milepæl
        pv_g = 0.0
        for w in wbs_list:
            p = PAKKE_BY_WBS[w]
            if måned < p.start_mnd:
                continue
            elif måned >= p.slutt_mnd:
                pv_g += p.bac
            else:
                pv_g += p.bac * (måned - p.start_mnd + 1) / p.varighet
        spi_g = ev_g / pv_g if pv_g > 0 else 0
        etc_g = bac_g - ev_g
        eac_g = bac_g / cpi_g if cpi_g > 0 else bac_g
        vac_g = bac_g - eac_g

        # Status
        pct_g = (ev_g / bac_g * 100) if bac_g > 0 else 0
        if pct_g >= 99.5:
            status = "Ferdig"
        elif pct_g == 0:
            status = "Ikke startet"
        elif spi_g < 0.9 or cpi_g < 0.9:
            status = "Avvik — følges opp"
        else:
            status = "På plan"

        row = t.rows[i]
        set_cell(row.cells[0], f"{mid}. {mnavn}")
        set_cell(row.cells[1], kr(bac_g))
        set_cell(row.cells[2], kr(ev_g))
        set_cell(row.cells[3], kr(ac_g))
        set_cell(row.cells[4], f"{cpi_g:.3f}" if cpi_g > 0 else "—")
        set_cell(row.cells[5], f"{spi_g:.3f}" if spi_g > 0 else "—")
        set_cell(row.cells[6], kr(etc_g))
        set_cell(row.cells[7], kr(vac_g))
        set_cell(row.cells[8], status)


# =============================================================================
# Tabell 4 — Avvik i perioden
# =============================================================================
def trim_table_rows(t, beholdt_antall: int) -> None:
    """Fjern rader etter `beholdt_antall` (inkludert header)."""
    rows_to_remove = list(t.rows[beholdt_antall:])
    for r in rows_to_remove:
        r._element.getparent().remove(r._element)


def fyll_avvik(doc, måned: int) -> None:
    """List avvik (vesentlige hendelser) som inntraff i denne måneden."""
    t = doc.tables[4]
    avvik_for_mnd = [h for h in HENDELSER if h.måned == måned]

    if not avvik_for_mnd:
        ensure_rows(t, 2)
        row = t.rows[1]
        set_cell(row.cells[0], "—")
        set_cell(row.cells[1], "Ingen vesentlige avvik registrert i perioden")
        set_cell(row.cells[2], "—")
        set_cell(row.cells[3], "—")
        set_cell(row.cells[4], "—")
        set_cell(row.cells[5], "—")
        set_cell(row.cells[6], "Lukket")
        trim_table_rows(t, 2)
        return

    ensure_rows(t, len(avvik_for_mnd) + 1)
    for i, h in enumerate(avvik_for_mnd, start=1):
        row = t.rows[i]
        cr_suffix = f" / {h.cr_id}" if h.cr_id else ""
        avvik_id = f"ISS-{h.måned:02d}-{i:02d} ({h.risiko_id}{cr_suffix})"
        konsekvens_text = (
            f"Pakke {h.pakke_påvirket}: "
            f"+{h.kostnad_mnok:.1f} MNOK risikoreserve, "
            f"+{h.tidsforskyvning_uker:.1f} uker tidsforskyvning"
        )
        set_cell(row.cells[0], avvik_id)
        set_cell(row.cells[1], h.tittel)
        set_cell(row.cells[2], konsekvens_text)
        set_cell(row.cells[3], h.beslutning)
        set_cell(row.cells[4], h.ansvarlig)
        set_cell(row.cells[5], h.frist)
        set_cell(row.cells[6], "Lukket" if h.cr_id or h.kostnad_mnok > 0 else "Overvåkes")

    trim_table_rows(t, len(avvik_for_mnd) + 1)


# =============================================================================
# Tabell 5 — Risikobudsjett
# =============================================================================
def fyll_risikobudsjett(doc, måned: int) -> None:
    evm = beregn_evm(måned)
    t = doc.tables[5]

    rr_rest = RISIKORESERVE - evm.risikoreserve_brukt_kum
    tb_rest = TIDSBUFFER_UKER - evm.risikoreserve_brukt_uker_kum

    # Row 1: Godkjent reserve (kr) | Brukt til nå (kr)
    set_cell(t.rows[1].cells[1], kr(RISIKORESERVE))
    set_cell(t.rows[1].cells[3], kr(evm.risikoreserve_brukt_kum))
    # Row 2: Gjenværende reserve (kr) | Åpne høye risikoer (antall)
    set_cell(t.rows[2].cells[1], kr(rr_rest))
    åpne_risiko = sum(1 for h in HENDELSER if h.måned > måned)
    set_cell(t.rows[2].cells[3], f"{åpne_risiko} (av registrerte)")
    # Row 3: Godkjent tidsbuffer (uker) | Brukt til nå (uker)
    set_cell(t.rows[3].cells[0], "Godkjent tidsbuffer (uker)")
    set_cell(t.rows[3].cells[1], f"{TIDSBUFFER_UKER} uker")
    set_cell(t.rows[3].cells[2], "Brukt tidsbuffer til nå (uker)")
    set_cell(t.rows[3].cells[3], f"{evm.risikoreserve_brukt_uker_kum:.1f} uker")
    # Row 4: Gjenværende tidsbuffer | Risikobudsjett som % av BAC
    set_cell(t.rows[4].cells[0], "Gjenværende tidsbuffer (uker)")
    set_cell(t.rows[4].cells[1], f"{tb_rest:.1f} uker")
    set_cell(t.rows[4].cells[2], "Risikoreserve som % av BAC")
    set_cell(t.rows[4].cells[3], f"{RISIKORESERVE/BAC*100:.1f} %")


# =============================================================================
# Tabell 6 — Risikoer
# =============================================================================
def fyll_risikoer(doc, måned: int) -> None:
    """List risikoer som er aktive eller nylig håndtert."""
    t = doc.tables[6]
    # Bruk alle 3 risikoene som vi har registrert hendelser for
    risiko_data = []
    for h in HENDELSER:
        if h.måned <= måned:
            risiko_data.append(h)

    # Hvis ingen aktive risikoer, sett en placeholder
    if not risiko_data:
        ensure_rows(t, 2)
        row = t.rows[1]
        set_cell(row.cells[0], "—")
        set_cell(row.cells[1], "Ingen risikoer har utløst respons så langt")
        for ci in range(2, 8):
            set_cell(row.cells[ci], "—")
        trim_table_rows(t, 2)
        return

    ensure_rows(t, len(risiko_data) + 1)
    for i, h in enumerate(risiko_data, start=1):
        row = t.rows[i]
        if h.måned == måned:
            status = "Aktiv — håndtert i perioden"
        elif h.kostnad_mnok > 0 or h.cr_id:
            status = "Lukket (oppfølging avsluttet)"
        else:
            status = "Lukket"
        konsekvens_niva = "Høy" if h.kostnad_mnok >= 5 else "Middels"

        cr_note = f" — endring {h.cr_id}" if h.cr_id else ""
        set_cell(row.cells[0], h.risiko_id)
        set_cell(row.cells[1], h.tittel + cr_note)
        set_cell(row.cells[2], "Realisert")
        set_cell(row.cells[3], konsekvens_niva)
        set_cell(row.cells[4], h.beslutning[:120] + ("…" if len(h.beslutning) > 120 else ""))
        set_cell(row.cells[5], kr(h.kostnad_mnok))
        set_cell(row.cells[6], f"{h.tidsforskyvning_uker:.1f} uker")
        set_cell(row.cells[7], status)

    trim_table_rows(t, len(risiko_data) + 1)


# =============================================================================
# Tabell 7 — S-kurve datatabell
# =============================================================================
def fyll_s_kurve_data(doc, måned: int) -> None:
    """Kumulative PV/EV/AC/CPI/SPI per måned for hele prosjektforløpet."""
    t = doc.tables[7]
    # Vi trenger 17 rader (1 header + 16 måneder)
    ensure_rows(t, 17)

    for m in range(1, 17):
        evm = beregn_evm(m)
        navn, _ = MÅNEDER[m]
        row = t.rows[m]
        marker = " ← Statusdato" if m == måned else ""
        set_cell(row.cells[0], f"Mnd {m}{marker}")
        set_cell(row.cells[1], navn)
        set_cell(row.cells[2], kr(evm.pv_kum))
        if m <= måned:
            set_cell(row.cells[3], kr(evm.ev_kum))
            set_cell(row.cells[4], kr(evm.ac_kum))
            set_cell(row.cells[5], f"{evm.cpi:.3f}")
            set_cell(row.cells[6], f"{evm.spi:.3f}")
        else:
            set_cell(row.cells[3], "—")
            set_cell(row.cells[4], "—")
            set_cell(row.cells[5], "—")
            set_cell(row.cells[6], "—")


# =============================================================================
# Figurer (S-kurve + Gantt-plassholder)
# =============================================================================
def insert_s_kurve_after_placeholder(doc, måned: int) -> None:
    """Sett inn S-kurve-bilde etter 'Figur 2'-plassholderen."""
    s_path = S_KURVER / f"s_kurve_mnd_{måned:02d}.png"
    if not s_path.exists():
        print(f"  ADVARSEL: S-kurve mangler for mnd {måned}: {s_path}")
        return
    for p in doc.paragraphs:
        if "Figur 2" in p.text:
            # Tøm placeholder-paragrafen og sett inn bilde der
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            run = p.add_run()
            run.add_picture(str(s_path), width=Cm(16))
            return


# =============================================================================
# Analytisk tekst per seksjon (erstatter mal-tekstene)
# =============================================================================
def erstatt_paragraf(doc, match_tekst: str, ny_tekst: str) -> bool:
    """Finn en paragraf hvis tekst inneholder match_tekst og erstatt den."""
    for p in doc.paragraphs:
        if match_tekst in p.text:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            run = p.add_run(ny_tekst)
            return True
    return False


def beskriv_cpi(cpi: float) -> str:
    if cpi >= 1.02:
        return "betydelig under budsjett (gunstig kostnadsavvik)"
    if cpi >= 1.0:
        return "innenfor budsjett"
    if cpi >= 0.95:
        return "marginalt over budsjett (akseptabelt)"
    if cpi >= 0.9:
        return "moderat over budsjett (krever oppfølging)"
    return "vesentlig over budsjett (krever korrigerende tiltak)"


def beskriv_spi(spi: float) -> str:
    if spi >= 1.05:
        return "forut for plan"
    if spi >= 1.0:
        return "iht. plan"
    if spi >= 0.95:
        return "marginalt bak plan (innenfor toleranse)"
    if spi >= 0.9:
        return "moderat bak plan (følges opp tett)"
    return "vesentlig bak plan (krever korrigerende tiltak)"


def fjern_paragraf(doc, match_tekst: str) -> bool:
    """Fjern en paragraf hvis tekst inneholder match_tekst."""
    for p in doc.paragraphs:
        if match_tekst in p.text:
            p._element.getparent().remove(p._element)
            return True
    return False


def oppdater_footer_og_tittel(doc, måned: int) -> None:
    """Erstatt "Mal for..."-tekst i footer og tittel-paragraf."""
    navn, _ = MÅNEDER[måned]
    rapport_ident = f"Månedsrapport — Nye Hædda Barneskole — {navn} — Mnd {måned}/16"
    # Fjern første "Mal for"-paragraf i body
    fjern_paragraf(doc, "Mal for månedlig prosjektstyringsrapport")
    # Footer (i sections)
    for section in doc.sections:
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer is None:
                continue
            for p in footer.paragraphs:
                if "Mal for" in p.text:
                    for r in list(p.runs):
                        r._element.getparent().remove(r._element)
                    p.add_run(rapport_ident)


def fyll_analytiske_kommentarer(doc, måned: int) -> None:
    """Erstatter mal-teksten under hver heading med periodespesifikk analyse."""
    oppdater_footer_og_tittel(doc, måned)
    # Fjern den generelle mal-intro-teksten (bruksanvisning for mal-bruker)
    fjern_paragraf(doc, "Bruk denne malen sammen med tilhørende Excel-arbeidsbok")
    evm = beregn_evm(måned)
    navn, statusdato = MÅNEDER[måned]
    rr_pct = (evm.risikoreserve_brukt_kum / RISIKORESERVE) * 100
    tb_pct = (evm.risikoreserve_brukt_uker_kum / TIDSBUFFER_UKER) * 100
    avvik_for_mnd = [h for h in HENDELSER if h.måned == måned]

    # Seksjon 1 — Hoved-KPI-er
    seksjon1 = (
        f"Ved statusdato {statusdato} (slutten av {navn}) har prosjektet et opptjent verdi (EV) "
        f"på {evm.ev_kum:.1f} MNOK ({evm.pct_complete:.1f} % av BAC på {BAC:.0f} MNOK), "
        f"med en faktisk kostnad (AC) på {evm.ac_kum:.1f} MNOK og planlagt verdi (PV) på "
        f"{evm.pv_kum:.1f} MNOK. CPI på {evm.cpi:.3f} viser at prosjektet er "
        f"{beskriv_cpi(evm.cpi)}, mens SPI på {evm.spi:.3f} viser at prosjektet er "
        f"{beskriv_spi(evm.spi)}. Estimert sluttkost (EAC) er {evm.eac:.1f} MNOK, som gir "
        f"en variance at completion (VAC) på {evm.vac:+.1f} MNOK mot vedtatt BAC. "
        f"Overordnet status: {evm.status_rag}."
    )
    erstatt_paragraf(doc, "Oppsummer prosjektets nåværende status", seksjon1)

    # Seksjon 2 — Gantt-figur (innledende tekst)
    seksjon2 = (
        f"Figur 1 viser Baseline 1 (godkjent etter komprimering av 4.1 Råbygg iht. kommunestyrets "
        f"vedtak {SAKSNUMMER_NHB}) sammen med faktisk fremdrift pr. statusdato {statusdato}. "
        "Eksporteres fra MS Project og settes inn under (Gantt — Tracking-visning med "
        "baseline + faktisk-stolper)."
    )
    erstatt_paragraf(doc, "Sett inn et bilde eksportert fra MS Project", seksjon2)

    # Erstatte Figur 1-plassholder med tekst som ber bruker lime inn fra MS Project
    figur1_tekst = (
        f"[FIGUR 1 — settes inn manuelt: Tracking-Gantt fra MS Project, statusdato {statusdato}. "
        f"Vis Baseline 1-stolper sammen med faktisk fremdrift. Marker statusdato med dagens linje.]"
    )
    erstatt_paragraf(doc, "Figur 1. Erstatt denne plassholderen", figur1_tekst)

    # Seksjon 3 — Avvik
    if not avvik_for_mnd:
        seksjon3 = (
            "Ingen vesentlige avvik registrert i perioden. Aktive risikoer overvåkes iht. "
            "risikoplan, og bruk av risikoreserve og tidsbuffer er uendret denne måneden. "
            "Tidligere registrerte hendelser følges opp iht. fastsatte frister."
        )
    else:
        avvik_beskrivelser = []
        for h in avvik_for_mnd:
            tekst = (
                f"{h.risiko_id}"
                + (f" / {h.cr_id}" if h.cr_id else "")
                + f": {h.tittel}. "
                f"Konsekvens: {h.kostnad_mnok:.1f} MNOK risikoreserve, "
                f"{h.tidsforskyvning_uker:.1f} uker tidsforskyvning. Beslutning iht. møtereferat "
                f"og endringsdokument."
            )
            avvik_beskrivelser.append(tekst)
        antall = len(avvik_for_mnd)
        innledning = (
            f"{antall} vesentlig avvik registrert i perioden, håndtert iht. risikoplan."
            if antall == 1 else
            f"{antall} vesentlige avvik registrert i perioden, alle håndtert iht. risikoplan."
        )
        seksjon3 = (
            f"{innledning} "
            f"{' '.join(avvik_beskrivelser)} Korrigerende tiltak er iverksatt, og oppfølging "
            f"skjer mot ansvarlige med fastsatte frister."
        )
    erstatt_paragraf(doc, "List bare vesentlige avvik", seksjon3)

    # Seksjon 4 — Risikoer
    aktive_hittil = [h for h in HENDELSER if h.måned <= måned]
    seksjon4 = (
        f"Brukt risikoreserve hittil: {evm.risikoreserve_brukt_kum:.1f} MNOK av godkjent "
        f"{RISIKORESERVE:.0f} MNOK ({rr_pct:.1f} %). "
        f"Brukt tidsbuffer hittil: {evm.risikoreserve_brukt_uker_kum:.1f} uker av godkjent "
        f"{TIDSBUFFER_UKER} uker ({tb_pct:.1f} %). "
        f"Antall hendelser med registrert respons hittil: {len(aktive_hittil)} "
        f"({', '.join(h.risiko_id for h in aktive_hittil) if aktive_hittil else 'ingen'}). "
    )
    if aktive_hittil:
        seksjon4 += (
            "Restrisiko vurderes som håndterbar — gjenværende reserve og buffer dekker "
            "rimelig restexpose for de pakkene som fortsatt er aktive."
        )
    else:
        seksjon4 += "Restrisiko vurderes som lav på dette stadiet i prosjektet."
    erstatt_paragraf(doc, "Beskriv dagens risikobilde", seksjon4)

    # Seksjon 5 — S-kurve
    seksjon5 = (
        f"S-kurven nedenfor viser kumulativ PV, EV og AC for hele prosjektforløpet "
        f"(mnd 1–16), med statuslinje ved mnd {måned} ({navn}). Prosjektet er "
        f"{evm.pct_complete:.1f} % fullført målt i opptjent verdi. "
        f"CPI på {evm.cpi:.3f} indikerer {beskriv_cpi(evm.cpi)}, og SPI på {evm.spi:.3f} "
        f"indikerer at prosjektet er {beskriv_spi(evm.spi)}. "
    )
    if måned <= 3:
        seksjon5 += (
            "Den høye SPI tidlig i prosjektet skyldes at detaljprosjektering (2.1) ble fullført "
            "raskere enn lineær fordeling tilsier — en planlagt frontlasting av planleggingsfasen."
        )
    elif måned >= 14:
        seksjon5 += (
            "Prosjektet nærmer seg overlevering med både CPI og SPI på eller over 1.0. "
            "EAC konvergerer mot BAC, og sluttdato 15.05.2026 er innenfor rekkevidde."
        )
    erstatt_paragraf(doc, "Lim inn den kumulative PV-/EV-/AC-S-kurven", seksjon5)

    # Datatabell-innledning
    datatab = (
        "Datatabellen under viser kumulative prosjektverdier per måned for hele forløpet. "
        f"Statusdato er markert med ← Statusdato i raden for mnd {måned}. "
        "Fremtidige måneder viser bare planlagt PV; EV/AC/CPI/SPI fylles ut etter hvert som "
        "statusrapportene leveres."
    )
    erstatt_paragraf(doc, "Datatabellen under figuren skal vise", datatab)


# =============================================================================
# Hovedfunksjon
# =============================================================================
def lag_månedsrapport(måned: int) -> Path:
    navn, _ = MÅNEDER[måned]
    out = OUTDIR / f"Månedsrapport mnd {måned:02d} - {navn}.docx"
    shutil.copy(MAL, out)
    doc = Document(out)

    fyll_metadata(doc, måned)
    fyll_kpi_sammendrag(doc, måned)
    fyll_milepaeler(doc, måned)
    fyll_avvik(doc, måned)
    fyll_risikobudsjett(doc, måned)
    fyll_risikoer(doc, måned)
    fyll_s_kurve_data(doc, måned)
    fyll_analytiske_kommentarer(doc, måned)
    insert_s_kurve_after_placeholder(doc, måned)

    doc.save(out)
    return out


def main():
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("måneder", nargs="*", type=int,
                        help="Måned-numre å generere (1-16). Tomt = alle.")
    args = parser.parse_args()
    målmåneder = args.måneder or list(range(1, 17))
    print(f"Genererer månedsrapporter for: {målmåneder}")
    for m in målmåneder:
        out = lag_månedsrapport(m)
        print(f"  -> {out.name}")
    print("Ferdig.")


if __name__ == "__main__":
    main()
