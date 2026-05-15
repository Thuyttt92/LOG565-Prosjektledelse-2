"""
Oppdaterer sluttrapporten med faktiske sluttall fra fase 3.

Tar 04 - Avslutning/Sluttrapport - Nye Hædda barneskole - UTKAST.docx,
erstatter [plassholdere] med faktiske tall fra log565_master_data, og
lagrer som "Sluttrapport - Nye Hædda barneskole.docx".
"""
from __future__ import annotations
import sys
import shutil
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).parent))
from log565_master_data import (
    BAC, RISIKORESERVE, TIDSBUFFER_UKER, SAKSNUMMER_NHB,
    HENDELSER, beregn_evm,
)

sys.path.insert(0, str(Path(__file__).resolve().parent))
from paths import AVSLUTNING, MAL_SLUTTRAPPORT
# UTKAST-versjonen er slettet etter migrering. Bruker sluttrapport-mal hvis regenerering.
UTKAST = MAL_SLUTTRAPPORT
OUT = AVSLUTNING / "Sluttrapport - Nye Hædda barneskole.docx"

# Faktiske sluttall (mnd 16)
sluttall = beregn_evm(16)
ANTALL_REALISERTE = len([h for h in HENDELSER])  # 3
RISIKORESERVE_BRUKT = sluttall.risikoreserve_brukt_kum  # 11 MNOK
TIDSBUFFER_BRUKT = sluttall.risikoreserve_brukt_uker_kum  # 1.5 uker
CPI_SLUTT = sluttall.cpi  # 1.000
SPI_SLUTT = sluttall.spi  # 1.000


def set_cell_simple(cell, tekst: str) -> None:
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p0 = cell.paragraphs[0]
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)
    p0.add_run(tekst)


def erstatt_i_paragraf(p, gammelt: str, nytt: str) -> bool:
    """Erstatt en streng i en paragraf på tvers av runs."""
    full = p.text
    if gammelt not in full:
        return False
    ny_tekst = full.replace(gammelt, nytt)
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(ny_tekst)
    return True


def erstatt_i_dokument(doc, erstatninger: dict[str, str]) -> int:
    """Gå gjennom alle paragrafer og tabell-celler. Returner antall erstatninger."""
    n = 0
    for p in doc.paragraphs:
        for gammelt, nytt in erstatninger.items():
            if erstatt_i_paragraf(p, gammelt, nytt):
                n += 1
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for gammelt, nytt in erstatninger.items():
                        if erstatt_i_paragraf(p, gammelt, nytt):
                            n += 1
    return n


def main():
    shutil.copy(UTKAST, OUT)
    doc = Document(OUT)

    # Tekstuelle erstatninger
    erstatninger = {
        "[DATO]": "15. mai 2026",
        "[SUM mill kr]": f"{sluttall.ac_kum:.0f} mill kr",
        "[Y mill kr — fra Bårds spesifikasjon]": "50 mill kr (komprimering av 4.1 Råbygg)",
        "[Y mill kr]": "50 mill kr",
        "[Z mill kr av 50 mill]": f"{RISIKORESERVE_BRUKT:.0f} mill kr av 50 mill kr",
        "[SPI=X, beskriv om vi var foran/bak skjema, hvilke perioder som drev avvik]":
            f"SPI={SPI_SLUTT:.3f} ved prosjektslutt (på prikken iht. plan). Prosjektet "
            "var marginalt forut for plan i tidlige måneder (mnd 1–3 SPI≈1,17–1,25) på grunn "
            "av rask fullføring av detaljprosjektering, og holdt SPI rundt 1,0 gjennom råbygg- "
            "og innvendigfasen. Tidsbuffer brukt: 1,5 uker av godkjent 8 uker (R-07 brann hos "
            "vindusprodusent i mnd 11)",
        "[CPI-verdi]": f"{CPI_SLUTT:.3f}",
        "[N] realisert": f"{ANTALL_REALISERTE} realisert",
        "[X dager / Y mill kr av 171 dager / 50 mill kr]":
            f"{TIDSBUFFER_BRUKT:.1f} uker / {RISIKORESERVE_BRUKT:.0f} mill kr av 8 uker / 50 mill kr",
        "[beskrivelse av eventuell realisert kritisk/høy risiko]":
            "R-06 (regulatorisk endring DSB sprinkler/rømning, mnd 13, +5 mill kr) — håndtert "
            "som endringsforespørsel CR-001. R-05 (forurenset masse under 3.2 Riving, mnd 4, "
            "+6 mill kr fra risikoreserve) og R-07 (brann hos vindusprodusent, mnd 11, "
            "+1,5 uker fra tidsbuffer) er de andre realiserte hendelsene",
        "[aktivitet]": "4.1 Råbygg",
        "[Eventuelle scope-endringer underveis dokumenteres her — se også endringsdokument "
        "CR-001 for crashing-saken.]":
            "Crashing av 4.1 Råbygg (NHB-2026-15) endret tid og kostnad, men ikke scope. "
            "DSB-veileder (CR-001) ga marginal scope-utvidelse på 5.1 VVS (oppgradert "
            "sprinklerdekning og rømningsskilting) — kompensert innenfor risikoreserve",
        "[Levert iht. plan]": "59 krav levert, 32 arbeidspakker fullført til 100 %",
        "[Faktisk sluttdato]": "15. mai 2026 (innen frist)",
        "[Faktisk totalkostnad]": f"{sluttall.ac_kum:.0f} mill kr (på prikken med BAC)",
        "Crashing-merkostnad: [Y mill kr]":
            "Crashing-merkostnad: 50 mill kr (4.1 Råbygg), totalramme utvidet fra 700 → 800 mill kr",
        "Krevde crashing av [aktivitet] for å nå frist":
            "Crashing av 4.1 Råbygg krevd for å nå frist (NHB-2026-15-vedtak)",
        "[Antall mangler]": "0 kritiske mangler (BP3 godkjent, ferdigbefaring lukket i mnd 15–16)",
        "[Liste over anbefalte tiltak etter prosjektets avslutning, typisk inkludert "
        "oppfølgingsaktiviteter, gjenstående forbedringstiltak og prioriterte saker fra "
        "bruker- eller sluttevalueringer.]":
            "1) Gevinstrealisering bør måles 12–24 mnd etter ibruktagelse — anbefales fulgt opp av "
            "driftsorganisasjonen mot business case (NNV +109 mill kr, BCR 1,16). "
            "2) FDV-dokumentasjon (5.5 IKT-sikkerhet, 5.4 SD-anlegg) bør gjennomgås årlig de "
            "første tre årene for å sikre at oppdaterte sprinkler-/rømningsløsninger (CR-001) "
            "fortsatt etterleves. "
            "3) Læringspunkter fra schedule crashing-prosessen (NHB-2026-15) tas inn i "
            "kommunens prosjektmal som case-eksempel for endringsstyring. "
            "4) Brukeropplæring (8.4) bør repeteres etter 6 mnd drift for å sikre at "
            "drift- og personalstaben behersker SD-anlegg og adgangskontroll.",
    }

    n = erstatt_i_dokument(doc, erstatninger)
    print(f"Erstattet {n} plassholdere.")

    # Oppdater status-tabellen (Tabell 0)
    t0 = doc.tables[0]
    set_cell_simple(t0.rows[0].cells[1], "15. mai 2026")
    set_cell_simple(t0.rows[1].cells[1], "1.0 — endelig sluttrapport")
    set_cell_simple(t0.rows[3].cells[1], "FERDIG")

    # Oppdater oppsummeringstabellen (Tabell 1)
    t1 = doc.tables[1]
    # Row 1: Omfang | 59 krav, 116 leveranser | (faktisk) | Crashing endret ikke scope
    set_cell_simple(t1.rows[1].cells[2], "59 krav levert, 32 arbeidspakker × 100 %")
    set_cell_simple(t1.rows[1].cells[3], "Crashing endret ikke scope. CR-001 ga marginal utvidelse på 5.1 VVS")
    # Row 2: Tid
    set_cell_simple(t1.rows[2].cells[2], "15. mai 2026 (på datoen)")
    set_cell_simple(t1.rows[2].cells[3], "Levert innen frist. 1,5 uker tidsbuffer brukt (av 8)")
    # Row 3: Kost
    set_cell_simple(t1.rows[3].cells[2], f"{sluttall.ac_kum:.0f} mill kr (BAC = 800)")
    set_cell_simple(t1.rows[3].cells[3],
                    "Rammeutvidelse 700→800 mill kr (NHB-2026-15). 11 mill kr risikoreserve brukt (22 %)")
    # Row 4: Kvalitet
    set_cell_simple(t1.rows[4].cells[2], "0 kritiske mangler, FDV-dokumentasjon levert")
    set_cell_simple(t1.rows[4].cells[3], "Brukstillatelse innvilget. K-001 og K-002 oppfylt")

    # Fjern "Status for utkastet"-vedlegget på slutten (den er bare for utkast-versjonen)
    # Finn paragrafer som hører til "Vedlegg: Status for utkastet" og fjern
    fjern_fra = False
    paragrafer_å_fjerne = []
    for p in doc.paragraphs:
        if "Vedlegg: Status for utkastet" in p.text:
            fjern_fra = True
        if fjern_fra:
            paragrafer_å_fjerne.append(p)
    for p in paragrafer_å_fjerne:
        p._element.getparent().remove(p._element)
    print(f"Fjernet {len(paragrafer_å_fjerne)} paragrafer som tilhørte utkast-vedlegget.")

    doc.save(OUT)
    print(f"Lagret: {OUT.name}")


if __name__ == "__main__":
    main()
